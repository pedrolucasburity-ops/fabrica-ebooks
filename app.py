import streamlit as st
import openai
import json
import io
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# --- 1. CONFIGURAÇÃO VISUAL ---
st.set_page_config(page_title="Infinity Factory 9.0 (Auto-Publisher)", layout="wide", page_icon="🚀")

def carregar_css():
    st.markdown("""
    <style>
        .stApp { background-color: #0E1117; color: #fff; }
        .stButton>button { background: linear-gradient(90deg, #1CB5E0 0%, #000851 100%); color: white; border: none; padding: 12px; width: 100%; font-weight: bold; border-radius: 8px; }
        [data-testid="stSidebar"] { background-color: #111; }
        h1, h2, h3 { color: white !important; }
        .stError { background-color: #500; color: #fcc; }
        .stSuccess { background-color: #050; color: #cfc; }
    </style>""", unsafe_allow_html=True)

# --- 2. LOGIN ---
def check_password():
    if "authenticated" not in st.session_state: st.session_state.authenticated = False
    if st.session_state.authenticated: return True
    col1, col2, col3 = st.columns([1,2,1])
    with col2:
        st.title("🔒 Login Factory 9.0")
        senha = st.text_input("Senha", type="password")
        if st.button("Entrar"):
            if senha == "admin": st.session_state.authenticated = True; st.rerun()
            else: st.error("Senha incorreta")
    return False

# --- 3. CÉREBRO DA IA (PROMPTS AVANÇADOS) ---
def get_client(api_key): 
    return openai.OpenAI(api_key=api_key, base_url="https://api.groq.com/openai/v1")

def gerar_sumario_vendedor(client, tema, publico, modelo):
    """Gera apenas títulos magnéticos, sem numeração chata."""
    prompt = f"""
    Crie um sumário de 5 Capítulos para um E-book Best-Seller.
    Tema: {tema}
    Público: {publico}
    
    REGRAS DE OURO PARA OS TÍTULOS:
    1. PROIBIDO títulos genéricos (Ex: "Introdução", "Conclusão").
    2. Use títulos de Copywriting que gerem curiosidade e promessa (Ex: "A Máquina de Renda", "O Segredo dos Bancos").
    3. Retorne APENAS a lista dos 5 títulos, um por linha. Nada mais.
    """
    try:
        resp = client.chat.completions.create(
            model=modelo,
            messages=[{"role": "user", "content": prompt}],
            temperature=0.8
        )
        return resp.choices[0].message.content
    except Exception as e:
        st.error(f"Erro Sumário: {e}")
        return None

def escrever_capitulo(client, titulo_atual, tema, publico, capitulos_anteriores, modelo):
    """Escreve o capítulo sabendo o que já foi dito para não repetir."""
    
    system_prompt = f"""
    Você é o autor de um Best-Seller sobre {tema}.
    Seu estilo é: DIRETO, PRÁTICO E ENGAJADOR.
    
    CONTEXTO DO LIVRO:
    - Já escrevemos sobre: {capitulos_anteriores}
    - Capítulo atual: {titulo_atual}
    
    REGRAS ESTRITAS (ANTI-REPETIÇÃO):
    1. NÃO defina conceitos básicos que já foram explicados antes. Vá direto ao ponto avançado.
    2. NÃO comece com "Neste capítulo vamos ver...". Comece com uma história, um dado impactante ou uma pergunta.
    3. Use Markdown: ## para subtítulos e **negrito** nas frases de impacto.
    4. Escreva parágrafos curtos (máximo 4 linhas) para leitura fácil.
    5. Termine com uma "Dica de Mestre" prática.
    """
    
    try:
        resp = client.chat.completions.create(
            model=modelo,
            messages=[{"role": "system", "content": system_prompt}, 
                      {"role": "user", "content": f"Escreva o capítulo '{titulo_atual}' agora. Mínimo 800 palavras."}],
            temperature=0.7
        )
        return resp.choices[0].message.content
    except Exception as e:
        return None

def gerar_bonus_pratico(client, tema, modelo):
    """Gera o passo a passo final obrigatoriamente."""
    prompt = f"""
    Escreva um BÔNUS FINAL para o livro {tema}.
    Título: PLANO DE AÇÃO DE 7 DIAS.
    
    Conteúdo: Crie um checklist passo a passo, dia por dia (Dia 1, Dia 2...), do que a pessoa tem que fazer na prática (clicar, abrir conta, comprar) para ter resultado.
    Sem teoria, apenas prática.
    """
    try:
        resp = client.chat.completions.create(
            model=modelo,
            messages=[{"role": "user", "content": prompt}],
            temperature=0.6
        )
        return resp.choices[0].message.content
    except: return None

# --- 4. GERADOR DE WORD ---
def criar_word(tema, publico, conteudo_raw, capa_file):
    doc = Document()
    
    # Capa
    if capa_file:
        try:
            doc.add_picture(capa_file, width=Inches(6))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        except: pass
        
    doc.add_paragraph("\n")
    t = doc.add_heading(tema.upper(), 0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Manual Prático para {publico}").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()
    
    # Processamento do Conteúdo
    if conteudo_raw:
        for linha in conteudo_raw.split('\n'):
            linha = linha.strip()
            if not linha: continue
            
            if linha.startswith('# '): 
                doc.add_heading(linha.replace('# ', ''), 1)
            elif linha.startswith('## '): 
                doc.add_heading(linha.replace('## ', ''), 2)
            elif '---' in linha: 
                doc.add_page_break()
            else: 
                p = doc.add_paragraph()
                parts = linha.split('**')
                for i, pt in enumerate(parts):
                    r = p.add_run(pt)
                    if i % 2 == 1: 
                        r.bold = True
                        r.font.color.rgb = RGBColor(0, 100, 0) # Verde escuro para destaque
                p.paragraph_format.space_after = Pt(12)
    
    b = io.BytesIO()
    doc.save(b)
    b.seek(0)
    return b

# --- APP ---
if "dados" not in st.session_state: st.session_state.dados = {"tema": "", "publico": "", "sumario_lista": [], "conteudo": ""}
carregar_css()

if check_password():
    st.sidebar.title("Factory 9.0 🚀")
    api_key = st.sidebar.text_input("API Key Groq", type="password")
    modelo = st.sidebar.selectbox("Modelo IA", ["llama-3.3-70b-versatile", "mixtral-8x7b-32768", "llama3-70b-8192"])

    if not api_key: st.stop()
    client = get_client(api_key)

    t1, t2, t3 = st.tabs(["1. Estrutura Viral", "2. Escrita Inteligente", "3. Download Pronto"])
    
    with t1:
        c1, c2 = st.columns(2)
        with c1:
            st.session_state.dados["tema"] = st.text_input("Tema", value="Guia de Dividendos")
            st.session_state.dados["publico"] = st.text_input("Público", value="Iniciantes com R$100")
        with c2:
            if st.button("Gerar Títulos Magnéticos"):
                res = gerar_sumario_vendedor(client, st.session_state.dados["tema"], st.session_state.dados["publico"], modelo)
                if res:
                    # Limpa e cria lista
                    lista = [l.strip() for l in res.split('\n') if l.strip() and not l.startswith('Aqui')]
                    st.session_state.dados["sumario_lista"] = lista
                    st.rerun()
        
        if st.session_state.dados["sumario_lista"]:
            st.success("Títulos Gerados (Pode editar se quiser):")
            for i, tit in enumerate(st.session_state.dados["sumario_lista"]):
                st.session_state.dados["sumario_lista"][i] = st.text_input(f"Capítulo {i+1}", value=tit, key=f"cap_{i}")

    with t2:
        if st.button("🚀 ESCREVER E-BOOK COMPLETO (AUTO-PUBLISHER)"):
            if not st.session_state.dados["sumario_lista"]:
                st.error("Gere os títulos na aba 1 primeiro!")
            else:
                bar = st.progress(0)
                st.session_state.dados["conteudo"] = ""
                titulos_usados = []
                total = len(st.session_state.dados["sumario_lista"]) + 1 # +1 do bonus

                # Loop dos Capítulos Normais
                for i, titulo in enumerate(st.session_state.dados["sumario_lista"]):
                    with st.spinner(f"Escrevendo: {titulo}... (Evitando repetições)"):
                        # O segredo: Passamos a lista do que já foi escrito
                        txt = escrever_capitulo(client, titulo, st.session_state.dados["tema"], st.session_state.dados["publico"], titulos_usados, modelo)
                        
                        if txt:
                            st.session_state.dados["conteudo"] += f"# {titulo}\n\n{txt}\n\n---\n"
                            titulos_usados.append(titulo)
                        else:
                            st.error(f"Erro no cap {i+1}")
                    bar.progress((i+1)/total)

                # Loop do Bônus Prático
                with st.spinner("Criando Plano de Ação Prático..."):
                    bonus = gerar_bonus_pratico(client, st.session_state.dados["tema"], modelo)
                    if bonus:
                        st.session_state.dados["conteudo"] += f"# BÔNUS: PLANO DE AÇÃO\n\n{bonus}\n"
                    bar.progress(1.0)
                
                st.balloons()
                st.success("Livro finalizado com sucesso! Vá para a aba 3.")

    with t3:
        st.header("Seu Produto Final")
        capa = st.file_uploader("Capa", type=['png', 'jpg'])
        if st.session_state.dados["conteudo"]:
            if st.button("Gerar Arquivo DOCX"):
                docx = criar_word(st.session_state.dados["tema"], st.session_state.dados["publico"], st.session_state.dados["conteudo"], capa)
                st.download_button("BAIXAR WORD PRONTO", docx, "Ebook_Vendedor.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
